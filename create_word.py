from utils import CheckOutliers
from utils.check_random import CheckRandom
from docx import Document
from docx.shared import Inches
import pandas as pd
import numpy as np
from scipy.stats import chi2
from utils import CheckCorell, k_means_clustering, assign_clusters, Autocorrel
from docx.enum.table import WD_TABLE_ALIGNMENT


def export_pirs_report(pirs_res: pd.DataFrame, doc: Document):
    """
    Добавляет в doc главу с результатами χ²‑теста Пирсона.
    pirs_res: DataFrame с колонками ['column','chi2_stat','p_value','df','note'].
    """
    doc.add_heading("Глава 1. χ²‑тест Пирсона на нормальность", level=1)
    for row in pirs_res.to_dict('records'):
        doc.add_heading(f"Столбец: {row['column']}", level=2)
        doc.add_paragraph(f"χ²‑статистика: {row['chi2_stat']}")
        doc.add_paragraph(f"Степени свободы: {row['df']}")
        doc.add_paragraph(f"p‑значение: {row['p_value']}")
        if row['note']:
            doc.add_paragraph(f"⚠️ Примечание: {row['note']}", style='IntenseQuote')
        conclusion = ("Нет оснований отвергнуть нормальность."
                      if row['p_value'] >= 0.05 else
                      "Отвергаем нормальность.")
        doc.add_paragraph(f"Вывод: {conclusion}", style='Quote')
        doc.add_paragraph()


def export_median_report(median_res: pd.DataFrame, doc: Document):
    """
    Добавляет в doc главу с результатами теста серий относительно медианы.
    median_res: DataFrame с колонками ['column','median','num_runs','expected_runs','conclusion'].
    """
    doc.add_heading("Глава 2. Тест на случайность методом серий", level=1)
    for row in median_res.to_dict('records'):
        doc.add_heading(f"Столбец: {row['column']}", level=2)
        doc.add_paragraph(f"Медиана: {row['median']}")
        doc.add_paragraph(f"Число серий: {row['num_runs']}")
        doc.add_paragraph(f"Ожидаемое число серий: {row['expected_runs']}")

        doc.add_paragraph(f"Максимальная серия: {row['max_run_length']}")
        doc.add_paragraph(f"Ожидаемая серия: {row['b crit']}")
        doc.add_paragraph(f"Вывод: Данные {row['conclusion']}", style='Quote')
        doc.add_paragraph()


def export_histograms_report(doc: Document, filename: str = "histograms.png"):
    """
    Добавляет в doc главу с гистограммами, используя build_hist.
    """
    # Построение и сохранение общей гистограммы
    doc.add_heading("Глава 3. Гистограммы с нормальной кривой", level=1)
    doc.add_picture(filename, width=Inches(6))
    doc.add_paragraph()


def export_outliers_report(outliers_res: pd.DataFrame, doc: Document):
    """
    Добавляет в doc главу с анализом выбросов.
    """
    doc.add_heading("Глава 4. Анализ выбросов", level=1)
    for row in outliers_res.to_dict('records'):
        doc.add_heading(f"Столбец: {row['column']}", level=2)
        doc.add_paragraph(f"Минимально допустимое значение: {row['lower_bound']}")
        doc.add_paragraph(f"Максимально допустимое значение: {row['upper_bound']}")
        doc.add_paragraph(f"Количество выбросов: {row['num_outliers']}")
        doc.add_paragraph(f"Процент выбросов: {row['percent_outliers']}%")
        doc.add_paragraph(f"Значения выбросов: {row['outlier_values']}")
        conclusion = ("Выбросов немного." if row['percent_outliers'] < 5
                      else "Присутствует значительное количество выбросов.")
        doc.add_paragraph(f"Вывод: {conclusion}", style='Quote')
        doc.add_paragraph()
def export_sigmas_report(sigmas_res: pd.DataFrame, doc: Document):
    """
    Добавляет в doc главу с анализом выбросов по правилу 3 сигм.
    """
    doc.add_heading("Глава 5. Анализ выбросов по правилу 3 сигм", level=1)
    for row in sigmas_res.to_dict('records'):
        doc.add_heading(f"Столбец: {row['column']}", level=2)
        doc.add_paragraph(f"Среднее: {row['mean']}")
        doc.add_paragraph(f"Стандартное отклонение: {row['std']}")
        doc.add_paragraph(f"Допустимый диапазон: от {row['lower_bound']} до {row['upper_bound']}")
        doc.add_paragraph(f"Количество выбросов: {row['num_outliers']}")
        doc.add_paragraph(f"Процент выбросов: {row['percent_outliers']}%")
        doc.add_paragraph(f"Значения выбросов: {row['outlier_values']}")
        conclusion = ("Выбросов немного." if row['percent_outliers'] < 5
                      else "Обнаружено значительное количество выбросов.")
        doc.add_paragraph(f"Вывод: {conclusion}", style='Quote')
        doc.add_paragraph()

def export_grabbs_report(grabbs_res: pd.DataFrame, doc: Document):
    """
    Добавляет в doc главу с результатами теста Граббса.
    """
    doc.add_heading("Глава 6. Тест Граббса на выбросы", level=1)
    doc.add_paragraph("⚠ Перед применением теста Граббса убедитесь, что данные имеют нормальное распределение.\n")

    for row in grabbs_res.to_dict('records'):
        doc.add_heading(f"Столбец: {row['column']}", level=2)
        doc.add_paragraph(f"Среднее: {row['mean']}")
        doc.add_paragraph(f"Стандартное отклонение: {row['std']}")
        doc.add_paragraph(f"Статистика G: {row['G']}")
        doc.add_paragraph(f"Критическое значение G: {row['G_crit']}")
        if row['outlier_value'] != '':
            doc.add_paragraph(f"Обнаруженный выброс: {row['outlier_value']}")
        doc.add_paragraph(f"Вывод: {row['conclusion']}", style='Quote')
        doc.add_paragraph()

def export_correlation_report(corr_matrix: pd.DataFrame, r2_matrix: pd.DataFrame,
                              t_matrix: pd.DataFrame, significance_matrix: pd.DataFrame,
                              doc: Document, alpha: float = 0.05):
    """
    Добавляет в doc главу с анализом корреляции между числовыми признаками.
    """

    def add_matrix_table(title: str, matrix: pd.DataFrame, formatter=str):
        doc.add_heading(title, level=2)
        table = doc.add_table(rows=1 + len(matrix), cols=1 + len(matrix.columns))
        table.style = 'Table Grid'

        # Заголовки
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = ""
        for j, col in enumerate(matrix.columns):
            hdr_cells[j + 1].text = str(col)

        # Данные
        for i, row_name in enumerate(matrix.index):
            row_cells = table.rows[i + 1].cells
            row_cells[0].text = str(row_name)
            for j, col in enumerate(matrix.columns):
                val = matrix.loc[row_name, col]
                row_cells[j + 1].text = formatter(val)

    doc.add_heading("Глава 6. Анализ корреляции между числовыми признаками", level=1)

    add_matrix_table("Матрица коэффициентов корреляции (r):", corr_matrix,
                     lambda v: f"{v:.3f}")
    add_matrix_table("Матрица детерминации (R²):", r2_matrix,
                     lambda v: f"{v:.3f}")
    add_matrix_table("Матрица t-статистик:", t_matrix,
                     lambda v: f"{v:.3f}")
    add_matrix_table(f"Матрица значимости (t_op > t_critical, α = {alpha}):",
                     significance_matrix, lambda v: "ЗН" if v else "НЗ")




def export_regression_report(models: dict, doc: Document):
    """
    Добавляет в doc главу с результатами построения и анализа регрессионных моделей.
    models: словарь с параметрами моделей, ключ — (X, Y), значения — dict с параметрами ('k', 'b', 'σ_ост', 'η²', 'TSS', 'ESS', 'RSS')
    """
    doc.add_heading("Глава 7. Анализ качества моделей линейной регрессии", level=1)

    for (x_col), params in models.items():
        doc.add_heading(f"Модель: y = k * {x_col} + b", level=2)
        doc.add_paragraph(f"Коэффициент наклона (k): {params['k']:.4f}")
        doc.add_paragraph(f"Свободный член (b): {params['b']:.4f}")
        doc.add_paragraph(f"Среднеквадратичное отклонение остатков (σ_ост): {params['σ_ост']:.4f}")
        doc.add_paragraph(f"Коэффициент детерминации (η²): {params['η²']:.4f}")
        doc.add_paragraph(f"Полная сумма квадратов (TSS): {params['TSS']:.4f}")
        doc.add_paragraph(f"Объяснённая сумма квадратов (ESS): {params['ESS']:.4f}")
        doc.add_paragraph(f"Остаточная сумма квадратов (RSS): {params['RSS']:.4f}")

        tss = params['TSS']
        ess = params['ESS']
        rss = params['RSS']
        check = "Да" if np.isclose(tss, ess + rss, atol=1e-4) else "Нет"
        doc.add_paragraph(f"Проверка: TSS ≈ ESS + RSS → {tss:.4f} ≈ {ess + rss:.4f} → Корректность: {check}")
        doc.add_paragraph()


def export_clustering_report(report_data: dict, doc: Document):
    """
    Добавляет в doc главу с результатами кластеризации (локоть, силуэт, кластеры).
    cluster_results: словарь с ключами:
        - 'elbow_plot': путь к изображению локтя
        - 'silhouette_plot': путь к изображению силуэта
        - 'best_k': оптимальное k
        - 'cluster_stats': список словарей по каждому кластеру с {'count', 'means'}
    """

    metrics_image_path = "clustering_metrics.png"
    doc.add_heading("Глава 8. Кластеризация методом k-средних", level=1)

    doc.add_paragraph(
        "Метод k-средних применяется для разделения объектов на группы (кластеры) на основании схожести значений признаков. "
        "Для определения оптимального числа кластеров используются два основных показателя: инерция (метод локтя) и коэффициент силуэта.")

    # Вставляем график
    doc.add_heading("Методы выбора количества кластеров", level=2)
    doc.add_picture(metrics_image_path, width=Inches(6))
    doc.add_paragraph(
        "Слева — график метода локтя: инерция (сумма квадратов расстояний до центроидов) уменьшается при увеличении k. "
        "Излом на графике указывает на оптимальное число кластеров.\n"
        "Справа — коэффициент силуэта, который показывает, насколько хорошо объекты соответствуют своему кластеру по сравнению с соседними. "
        "Чем выше значение, тем качественнее кластеризация."
    )

    # Определяем лучшее k
    # Жестко задаём оптимальное k = 13
    fixed_k = 13
    best_result = None
    for res in report_data:
        if res['k'] == fixed_k:
            best_result = res
            break

    if best_result is None:
        doc.add_paragraph(f"Оптимальное число кластеров k={fixed_k} не найдено в данных.")
        return

    doc.add_heading(f"Оптимальное число кластеров: {fixed_k}", level=2)
    doc.add_paragraph(
        f"Значение коэффициента силуэта для k={fixed_k}: {best_result['silhouette']:.3f}."
    )

    # Перебираем все кластеры
    for cluster in best_result['clusters']:
        doc.add_heading(f"Кластер {cluster['cluster']} — {cluster['count']} объектов", level=3)
        for key, val in cluster['mean_values'].items():
            doc.add_paragraph(f"{key}: {val}", style='List Bullet')



def export_multiple_regression_report(model_params: dict, x1_col: str, x2_col: str, y_col: str, doc):
    """
    Экспорт отчёта по одной модели множественной регрессии.

    model_params: dict - словарь с параметрами модели (k1, k2, b, σ_ост и т.д.)
    x1_col, x2_col, y_col: str - имена колонок, чтобы красиво оформить отчет
    doc: Document - объект docx для записи
    """

    doc.add_heading("Глава 7.1. Анализ качества модели множественной линейной регрессии", level=1)

    doc.add_heading(
        f"Модель: {y_col} = {model_params['k1']:.4f} * {x1_col} + {model_params['k2']:.4f} * {x2_col} + {model_params['b']:.4f}",
        level=2
    )
    doc.add_paragraph(f"Коэффициент при {x1_col} (k1): {model_params['k1']:.4f}")
    doc.add_paragraph(f"Коэффициент при {x2_col} (k2): {model_params['k2']:.4f}")
    doc.add_paragraph(f"Свободный член (b): {model_params['b']:.4f}")
    doc.add_paragraph(f"Среднеквадратичное отклонение остатков (σ_ост): {model_params['σ_ост']:.4f}")
    doc.add_paragraph(f"Коэффициент детерминации (η²): {model_params['η²']:.4f}")
    doc.add_paragraph(f"Полная сумма квадратов (TSS): {model_params['TSS']:.4f}")
    doc.add_paragraph(f"Объяснённая сумма квадратов (ESS): {model_params['ESS']:.4f}")
    doc.add_paragraph(f"Остаточная сумма квадратов (RSS): {model_params['RSS']:.4f}")

    tss = model_params['TSS']
    ess = model_params['ESS']
    rss = model_params['RSS']
    check = "Да" if np.isclose(tss, ess + rss, atol=1e-4) else "Нет"
    doc.add_paragraph(f"Проверка: TSS ≈ ESS + RSS → {tss:.4f} ≈ {ess + rss:.4f} → Корректность: {check}")
    doc.add_paragraph()

def export_multiple_regression_report_3x(model_params: dict, x1_col: str, x2_col: str, x3_col: str, y_col: str, doc):
    """
    Экспорт отчёта по модели множественной линейной регрессии с тремя признаками.

    model_params: dict - словарь с параметрами модели (k1, k2, k3, b, σ_ост и т.д.)
    x1_col, x2_col, x3_col, y_col: str - имена колонок для оформления отчёта
    doc: Document - объект python-docx для записи отчёта
    """

    doc.add_heading("Глава 7.1. Анализ качества модели множественной линейной регрессии", level=1)

    doc.add_heading(
        f"Модель: {y_col} = "
        f"{model_params['k1']:.4f} * {x1_col} + "
        f"{model_params['k2']:.4f} * {x2_col} + "
        f"{model_params['k3']:.4f} * {x3_col} + "
        f"{model_params['b']:.4f}",
        level=2
    )
    doc.add_paragraph(f"Коэффициент при {x1_col} (k1): {model_params['k1']:.4f}")
    doc.add_paragraph(f"Коэффициент при {x2_col} (k2): {model_params['k2']:.4f}")
    doc.add_paragraph(f"Коэффициент при {x3_col} (k3): {model_params['k3']:.4f}")
    doc.add_paragraph(f"Свободный член (b): {model_params['b']:.4f}")
    doc.add_paragraph(f"Среднеквадратичное отклонение остатков (σ_ост): {model_params['σ_ост']:.4f}")
    doc.add_paragraph(f"Коэффициент детерминации (η²): {model_params['η²']:.4f}")
    doc.add_paragraph(f"Полная сумма квадратов (TSS): {model_params['TSS']:.4f}")
    doc.add_paragraph(f"Объяснённая сумма квадратов (ESS): {model_params['ESS']:.4f}")
    doc.add_paragraph(f"Остаточная сумма квадратов (RSS): {model_params['RSS']:.4f}")

    tss = model_params['TSS']
    ess = model_params['ESS']
    rss = model_params['RSS']
    check = "Да" if np.isclose(tss, ess + rss, atol=1e-4) else "Нет"
    doc.add_paragraph(f"Проверка: TSS ≈ ESS + RSS → {tss:.4f} ≈ {ess + rss:.4f} → Корректность: {check}")
    doc.add_paragraph()



from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT

def export_correlation_by_clusters(df_with_clusters: pd.DataFrame, doc: Document, alpha: float = 0.05):
    summary = []  # сюда собираем данные для итоговой таблицы

    for cluster in sorted(df_with_clusters['cluster'].unique()):
        cluster_df = df_with_clusters[df_with_clusters['cluster'] == cluster].drop(columns=['cluster'])

        cluster_size = len(cluster_df)  # размер кластера

        doc.add_heading(f"Глава 9.{cluster + 1}. Корреляционный анализ: кластер {cluster}", level=1)

        corr_matrix, r2_matrix, t_matrix, significance_matrix = CheckCorell.check_correlation(cluster_df, alpha)
        export_correlation_report(corr_matrix, r2_matrix, t_matrix, significance_matrix, doc, alpha)
        z = CheckCorell.select_features_for_y(cluster_df, "Y", significance_matrix)

        eta2 = None
        sigma_ost = None

        if len(z) == 2:
            models = CheckCorell.build_multiple_regression_model(cluster_df, z[0], z[1], 'Y')
            export_multiple_regression_report(models, z[0], z[1], "Y", doc)

            eta2 = models.get('η²') if 'η²' in models else None
            sigma_ost = models.get('σ_ост') if 'σ_ост' in models else None

        else:
            models = CheckCorell.build_regression_models(cluster_df, significance_matrix, "Y")
            export_regression_report(models, doc)

            if models:
                first_key = next(iter(models))
                eta2 = models[first_key].get('η²')
                sigma_ost = models[first_key].get('σ_ост')

        summary.append({
            'cluster': cluster,
            'size': cluster_size,
            'η²': eta2,
            'σ_ост': sigma_ost
        })

    # Итоговая таблица с размером кластера
    doc.add_heading("Итоговая таблица по кластерам: размер, точность модели и σ_ост", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light List Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Номер кластера'
    hdr_cells[1].text = 'Размер кластера'
    hdr_cells[2].text = 'Коэффициент детерминации (η²)'
    hdr_cells[3].text = 'Среднеквадратичное отклонение остатков (σ_ост)'

    for row_data in summary:
        row_cells = table.add_row().cells
        row_cells[0].text = str(row_data['cluster'] + 1)
        row_cells[1].text = str(row_data['size'])
        row_cells[2].text = f"{row_data['η²']:.4f}" if row_data['η²'] is not None else "—"
        row_cells[3].text = f"{row_data['σ_ост']:.4f}" if row_data['σ_ост'] is not None else "—"



def export_lags_report(corr_df: pd.DataFrame, r2_df: pd.DataFrame, significance_df: pd.DataFrame, doc: Document):
    """
    Экспортирует автокорреляции (r), R² и значимость в виде таблицы для каждого X.
    """
    doc.add_heading("Глава 10. Анализ автокорреляции по лагам", level=1)

    for x_var in corr_df.index:
        doc.add_heading(f"Переменная: {x_var}", level=2)

        # Подготовка таблицы: строки — лаги, колонки — r, R², Значимость
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        # Заголовки
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Лаг"
        hdr_cells[1].text = "Корреляция (r)"
        hdr_cells[2].text = "Детерминация (R²)"
        hdr_cells[3].text = "Значимо (t-test)"

        for lag_label in corr_df.columns:
            row = table.add_row().cells
            row[0].text = lag_label

            r = corr_df.loc[x_var, lag_label]
            r2 = r2_df.loc[x_var, lag_label]
            signif = significance_df.loc[x_var, lag_label]

            row[1].text = str(r) if pd.notna(r) else ""
            row[2].text = str(r2) if pd.notna(r2) else ""
            row[3].text = "Да" if signif else "Нет"

        doc.add_paragraph()


def create(df):

    cr = CheckRandom()
    co = CheckOutliers()
    # Сначала получаем результаты тестов:
    pirs_df = cr.check_pirs(df)
    median_df = cr.check_median(df)
    cr.build_hist(df)
    # Создаём документ и добавляем главы
    doc = Document()
    doc.add_heading("Статистический отчёт", level=0)
    out = co.check_quantiles(df)
    out2 = co.check_sigmas(df)
    out3 = co.check_grabbs(df)

    # Корреляционный анализ
    corr_matrix, r2_matrix, t_matrix, significance_matrix = CheckCorell.check_correlation(df)
    models = CheckCorell.build_regression_models(df, significance_matrix, "Y")
    export_pirs_report(pirs_df, doc)
    export_median_report(median_df, doc)
    export_histograms_report(doc)
    export_outliers_report(out, doc)
    export_sigmas_report(out2, doc)
    export_grabbs_report(out3, doc)

    export_correlation_report(corr_matrix, r2_matrix, t_matrix, significance_matrix, doc)
    export_regression_report(models, doc)

    clusters_res = k_means_clustering(df, [i for i in range(2, 15)])
    export_clustering_report(clusters_res, doc)

    df_clustered = assign_clusters(df, k=13)
    export_correlation_by_clusters(df_clustered, doc)
    # Сохраняем
    doc.save("stat_report.docx")
    print("Отчёт сохранён в stat_report.docx")



def create2task(df):

    cr = CheckRandom()
    co = CheckOutliers()
    # Сначала получаем результаты тестов:
    pirs_df = cr.check_pirs(df)
    median_df = cr.check_median(df)
    cr.build_hist(df)
    # Создаём документ и добавляем главы
    doc = Document()
    doc.add_heading("Статистический отчёт", level=0)
    out = co.check_quantiles(df)
    out2 = co.check_sigmas(df)
    out3 = co.check_grabbs(df)

    # Корреляционный анализ

    export_pirs_report(pirs_df, doc)
    export_median_report(median_df, doc)
    export_histograms_report(doc)
    export_outliers_report(out, doc)
    export_sigmas_report(out2, doc)
    export_grabbs_report(out3, doc)

    corr_df, r2_df, significance_df = Autocorrel.check_lags(df, lag_count=25, y_col="Y", alpha=0.05)
    export_lags_report(corr_df, r2_df, significance_df, doc)

    shifted_df = Autocorrel.shift_by_significant_lags(df, corr_df, significance_df, y_col="Y")
    corr_matrix, r2_matrix, t_matrix, significance_matrix = CheckCorell.check_correlation(shifted_df)

    export_correlation_report(corr_matrix, r2_matrix, t_matrix, significance_matrix, doc)
    z = CheckCorell.select_features_for_y(shifted_df, "Y", significance_matrix)
    models = CheckCorell.build_multiple_regression_model_3x(shifted_df, "x2", "x3", "x4", 'Y')
    export_multiple_regression_report_3x(models,"x2", "x3", "x4", 'Y', doc)

    # if len(z) == 2:
    #
    #     models = CheckCorell.build_multiple_regression_model(shifted_df, z[0], z[1], 'Y')
    #     export_multiple_regression_report(models, z[0], z[1], "Y", doc)
    # else:
    #     models = CheckCorell.build_regression_models(shifted_df, significance_matrix, "Y")
    #     export_regression_report(models, doc)

    doc.save("stat_report_2_task.docx")
    print("Отчёт сохранён в stat_report.docx")